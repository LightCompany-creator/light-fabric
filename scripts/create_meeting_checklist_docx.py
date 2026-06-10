from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("docs/checklist_dlya_soveshaniya.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_check_table(doc, title, rows):
    h = doc.add_heading(title, level=2)
    h.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    widths = [Cm(5.0), Cm(10.4), Cm(2.0)]
    headers = ["Что запросить", "Какие данные нужны", "Статус"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        set_cell_shading(cell, "1F4E79")
        set_cell_text(cell, text, bold=True, color="FFFFFF", size=9)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for left, right in rows:
        cells = table.add_row().cells
        values = [left, right, "□"]
        for idx, value in enumerate(values):
            cells[idx].width = widths[idx]
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_text(cells[idx], value, size=8.7 if idx != 2 else 12)
            if idx == 2:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(9.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(31, 78, 121)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Что нужно собрать от руководителя и начальников цехов")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Чек-лист для совещания по программе Light Fabric")
    sr.font.size = Pt(11)
    sr.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Цель: собрать данные, чтобы система корректно считала выпуск, зарплату, сырье, остатки и отчеты.").bold = True

    doc.add_heading("Минимум, который нужно забрать сегодня", level=1)
    add_bullets(
        doc,
        [
            "Работники по каждому цеху: ФИО, табельный номер, должность, активность, кто открывает смену.",
            "Операции по каждому цеху: единые названия операций без вариантов написания.",
            "Расценки: цех, артикул, операция, сумма, единица, дата начала действия.",
            "Артикулы и маршруты: через какие цеха проходит каждый артикул.",
            "Нормы расхода материалов по артикулам и операциям.",
            "Правила учета брака, простоя, переделки и закрытия смены.",
            "Какие отчеты руководитель хочет видеть ежедневно, еженедельно и за месяц.",
        ],
    )

    doc.add_heading("Общее требование ко всем цехам", level=1)
    add_check_table(
        doc,
        "Данные, которые нужны от каждого начальника цеха",
        [
            ("Работники", "ФИО, табельный номер, должность, активен/не активен, кто может открывать и закрывать смену."),
            ("Операции", "Список операций с точным названием для программы: например, литье, обрезка, упаковка без носка."),
            ("Артикулы", "Коды и названия артикулов, которые проходят через цех."),
            ("Расценки", "Сколько рублей за пару, кг, операцию, смену; к какому артикулу и операции относится ставка."),
            ("Выработка", "Как считать количество: пары, штуки, кг, партии, формы; кто вводит данные и когда."),
            ("Брак", "Как фиксируется, оплачивается или нет, списывает ли сырье, кто отвечает."),
            ("Передача", "Откуда и куда передается продукция, кто подтверждает, нужна ли незавершенка между цехами."),
            ("Материалы", "Какие материалы или комплектующие расходуются в цехе, в каких единицах и по каким нормам."),
            ("Простой", "Нужно ли фиксировать простой, причины, минуты, влияет ли простой на зарплату."),
            ("Закрытие смены", "Какие данные обязательны перед закрытием смены и кто отвечает за правильность."),
        ],
    )

    doc.add_page_break()
    add_check_table(
        doc,
        "Примеры уточнений по разным цехам",
        [
            ("Станки", "Полный список станков и названия для программы: KCLKA 1, KCLKA 2, четверка, шестёрка и т.д."),
            ("Виды отливки", "4-парка, 6-парка, 8-парка; какие артикулы на каких формах и станках."),
            ("Вес и сырье", "Для литья: норма веса пары, ЭВА/ПВХ/краситель/добавки и нормы расхода."),
            ("Упаковка", "Пакет, коробка, носок, этикетка, вкладыш; нормы списания."),
            ("Артикулы", "Какие артикулы идут с носком и без носка, какие требуют отдельной упаковки."),
            ("Крой", "Какие детали кроятся, единица учета: деталь, пара, комплект; нормы материала и отходы."),
            ("Швейка", "Операции: пристрочка, манжет, подклад, ремешок; расценки и артикула по операциям."),
            ("Клеевой", "Склейка, вставка, подготовка, сушка; клей, растворитель, нормы расхода."),
            ("Обшив/сборка", "Какие операции выполняются, что считается готовым результатом, какие комплектующие списывать."),
            ("Маркировка", "Штамп, наклейка, бирка, этикетка; расценка и материалы списания."),
            ("Склад", "Статусы: принято, на складе, отгружено; кто принимает готовую продукцию."),
            ("Ангар/сырье", "Список материалов, единицы, начальные остатки, приходы, партии сырья, кто подтверждает списание."),
        ],
    )

    doc.add_page_break()
    add_check_table(
        doc,
        "Технолог, бухгалтер, руководитель",
        [
            ("Технолог", "Артикулы, вес пары, нормы расхода, маршруты по цехам, операции и действующие расценки."),
            ("Бухгалтер", "Форма ведомости ЗП, экспорт в Excel/1С, округление, брак, простой, переработки, удержания."),
            ("Руководитель", "Какие KPI видеть: выпуск, брак, ФОТ, сырье, остатки, эффективность; периоды отчетов."),
            ("Права доступа", "Кто директор, бухгалтер, технолог, начальник цеха; кто может исправлять закрытую смену."),
            ("Закрытая смена", "Что считается правильно закрытой сменой и какие поля обязательны перед закрытием."),
        ],
    )

    doc.add_page_break()
    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.add_heading("Фраза для совещания", level=1)
    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Cm(0.5)
    quote.paragraph_format.right_indent = Cm(0.5)
    quote.paragraph_format.space_before = Pt(4)
    quote.paragraph_format.space_after = Pt(8)
    qr = quote.add_run(
        "Чтобы программа считала зарплату, выпуск и сырье без ручных исправлений, "
        "мне от каждого цеха нужна связка: работник -> операция -> артикул -> количество -> "
        "расценка -> материал/норма. Если эта связка будет заполнена, система начнет работать корректно."
    )
    qr.italic = True
    qr.font.size = Pt(10.5)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build()
